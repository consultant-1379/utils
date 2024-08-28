'''
Created on Sept 09, 2021
@author: zpunvai
'''

import sys
import docx
import json
from paramiko import SSHClient
from scp import SCPClient
from collections import OrderedDict


########command line arguments#########
release_note = sys.argv[1]
zid = sys.argv[2]
passwrd = sys.argv[3]

doc = docx.Document('/JUMP/export/jumpstart/misc/json_convert/'+release_note)
final_data = OrderedDict()
product_data = []
media_data = []
ticket_data = []


#######For Product Details Table#########
product_table = doc.tables[1]
keys = None
for i,row in enumerate(product_table.rows):
        text = (cell.text.strip() for cell in row.cells)
        if i == 0:
                keys = tuple(text)
                continue
        final_data = OrderedDict(zip(keys, text))


########For Ticket Details Table#########
ticket_table = doc.tables[3]
keys = None
for i, row in enumerate(ticket_table.rows):
	text = (cell.text.strip() for cell in row.cells)
	if i == 0:
		keys = tuple(text)
		continue
	row_data = OrderedDict(zip(keys, text))
	ticket_data.append(row_data)
final_data['SoftwareGateway Tickets'] = ticket_data


########For Software Deliverable Table#######
deliverable_table = doc.tables[4]
keys = None
for i, row in enumerate(deliverable_table.rows):
        if i == 0:
                continue

        text = (cell.text.strip() for cell in row.cells)
        text_value = tuple(text)

        is_value_same = text_value.count(text_value[0]) == len(text_value)
        if is_value_same:
                continue

        if i == 1:
                keys = tuple(text_value)
                continue

        row_data = OrderedDict(zip(keys, text_value))
        media_data.append(row_data)
final_data['Software Deliverables'] = media_data


#######Changing function designation for ENM Media#######
final_data["Software Deliverables"][18]["Function Designation"] = "NAS config"
final_data["Software Deliverables"][20]["Function Designation"] = "NAS 7.4 Media"
final_data["Software Deliverables"][21]["Function Designation"] = "Redhat Enterprise Linux 7.6 Install ISO"
final_data["Software Deliverables"][22]["Function Designation"] = "NAS config"


########Json File Name########
AOM = final_data['Product Number']
AOM = AOM.replace(" ","")
RState = final_data['R-State']
json_file = "eniq_release_"+AOM+RState+".json"
print(json_file)


########Converting into JSON object and writting into file########
json_object = json.dumps(final_data, indent = 4)
with open ("/JUMP/export/jumpstart/misc/json_convert/"+json_file, "w") as outfile: 
	outfile.write(json_object)


#######Copying JSON file to User's Machine#######
ssh = SSHClient()
ssh.load_system_host_keys()
ssh.connect('seliius26954.seli.gic.ericsson.se',username=zid,password=passwrd)
scp = SCPClient(ssh.get_transport())
scp.put(json_file,json_file)
scp.close()
